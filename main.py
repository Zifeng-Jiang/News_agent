import streamlit as st
from pitcher_agent import PitcherAgent
from scripter_agent import ScripterAgent
from langgraph.graph import StateGraph
from typing import List, TypedDict
from scraper_tools import run_news_scraper
from docx import Document
from io import BytesIO
from event_scraper import *
from openai import OpenAI
import json
import re

class NewsState(TypedDict):
    region: str
    news_list: List[dict]
    selected_news: dict

# 构建完整的LangGraph
main_graph = StateGraph(NewsState)

pitcher_agent = PitcherAgent()
main_graph.add_node("pitcher", pitcher_agent.run)
scripter_agent = ScripterAgent()
main_graph.add_node("scripter", scripter_agent.run)
main_graph.add_edge("pitcher", "scripter")

# 设置图的入口点和出口点
main_graph.set_entry_point("pitcher")
main_graph.set_finish_point("scripter")

# streamlit app page
st.set_page_config(page_icon="📰", page_title="SpaceNews Agents")
st.title('📰SpaceNews Agents🤖')

"""Hello 👋🏻 You can get space/satellite news from all over the world through SpaceNews Agents."""
st.write("China | Middle East | Africa | Central Asia | Southeast Asia | Latin America | AI | Launch | Commercial | Events")
btn = st.button("Start Collecting and Summarizing")
message_placeholder = st.empty()
if btn:
    message_placeholder.empty()  # Clear previous success or error messages
    with st.spinner('Collecting and summarizing news...'):
        best_news = []
        news_list = run_news_scraper()
        # for news in news_list:
            # 初始化所有可能的字段，避免 KeyError
            # news.setdefault('title', '')
            # news.setdefault('link', '')
            # news.setdefault('date', '')
            # news.setdefault('tag', '')
            # news.setdefault('summary', '')
            # news.setdefault('abstract', '')
            # news.setdefault('content', '')
            #news.setdefault('images', [])

        regions = ["china", "middle_east", "africa", "central_asia", "southeast_asia", "latin_america", "AI", "launch", "commercial"]
        for region in regions:
            if region not in ["AI", "launch", "commercial"]:
                # 初始化状态
                initial_state = {
                    "region": region,
                    "news_list": news_list,
                    "selected_news": {}
                }
            elif region == 'AI':
                initial_state = {
                    "region": region,
                    "news_list": [news for news in news_list if news.get("tag") == "AI"],
                    "selected_news": {}
                }
            elif region == 'launch':
                initial_state = {
                    "region": region,
                    "news_list": [news for news in news_list if news.get("tag") == "launch"],
                    "selected_news": {}
                }
            else:
                initial_state = {
                    "region": region,
                    "news_list": [news for news in news_list if news.get("tag") == "commercial"],
                    "selected_news": {}
                }
            # 运行LangGraph
            app = main_graph.compile()
            final_state = app.invoke(initial_state)
            selected_news = final_state["selected_news"]
            if 'title' in selected_news:
                existing_titles = {news["title"] for news in best_news}

                if selected_news["title"] not in existing_titles:
                    selected_news["region"] = region 
                    
                    if 'tag' not in selected_news or selected_news['tag'] == '':           
                        tag_list = ['AI', 'Civil', 'Commercial', 'Finance', 'Launch', 'Opinion', 'Manufacturing', 'Imagery and Sensing']
                        example = {'tag': 'AI'}

                        base_urls = ["https://api.chatanywhere.tech/v1", "https://api.chatanywhere.com.cn/v1"]
                        client = OpenAI(
                            api_key="YOUR_OPEN_AI_API_KEY", 
                            base_url=base_urls[0],
                        )
                        completion = client.chat.completions.create(
                            model = "gpt-3.5-turbo",
                                messages = [
                                    {"role": "system", "content": f"You are a news editor expert in select the most relevant tag for news articles.\
                                    Choose only one tag from the following list: {tag_list}. \
                                    Return the selected tag as a JSON object with the key 'tag'. For example, {example}. Do not reply other information" },
                                    {"role": "user", "content": f"The news title is: '{selected_news['title']}'. News content is: '{selected_news['content'][:2000]}'."}
                                ],
                                temperature = 0.6,
                            )
                        tag_response = completion.choices[0].message.content
                        print("tag response:", tag_response)

                        # 解析大模型的输出
                        pattern = r'\{[^{}]*\}'
                        match = re.search(pattern, tag_response)
                        if match:
                            tag_response_str = match.group(0)
                            
                            # 替换单引号为双引号
                            tag_response_str = tag_response_str.replace("'", '"')
                            
                            try:
                                tag_response_dict = json.loads(tag_response_str)
                            except json.JSONDecodeError as e:
                                print(f"JSONDecodeError: {e}")
                                tag_response_dict = {"tag": 'Unknown'}
                        else:
                            tag_response_dict = {"tag": 'Unknown'}
                            print("Warning: Could not parse tag for news title.")
                        selected_news['tag'] = tag_response_dict['tag']
                    best_news.append(selected_news)

        events = get_events()

        # 爬取所有best_news的图片，并创建Word文档
        doc = Document()
        for article in best_news:
            doc.add_heading(article.get('region', 'N/A'), level=1)
            doc.add_heading(article.get('title', 'N/A'), level=2)
            doc.add_paragraph(f"Link: {article.get('link', 'N/A')}")
            doc.add_paragraph(f"Date: {article.get('date', 'N/A')}")
            doc.add_paragraph(f"Tag: {article.get('tag', 'N/A')}")
            doc.add_paragraph(f"Summary: {article.get('summary', 'N/A')}")
            doc.add_paragraph("\n")  # 每两个新闻条目之间间隔一行

        doc.add_heading('Events', level=1)

        for event in events:
            doc.add_heading(event.get('title', 'N/A'), level=2)
            doc.add_paragraph(f"Link: {event.get('link', 'N/A')}")
            doc.add_paragraph(f"Date: {event.get('date', 'N/A')}")
            doc.add_paragraph(f"Address: {event.get('address', 'N/A')}")
            doc.add_paragraph(f"Summary: {event.get('summary', 'N/A')}")
            doc.add_paragraph("\n")  # 每两个events条目之间间隔一行

        # 保存Word文档到内存缓冲区
        news_word = BytesIO()
        doc.save(news_word)
        news_word.seek(0)

        message_placeholder.success("News summarization completed successfully!")
        st.download_button(
            label="Download News Summary as Word",
            data=news_word,
            file_name="news_summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
